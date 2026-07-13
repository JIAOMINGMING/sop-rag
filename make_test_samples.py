#!/usr/bin/env python3
"""Create deliberately-messy test documents in docs/ to exercise the extractors:
- a Chinese .docx with NO heading styles (bold paragraphs as headings)
- a .pdf with numbered sections
- a deviation-log .xlsx
"""
from pathlib import Path

import fitz
from docx import Document
from openpyxl import Workbook

DOCS = Path(__file__).parent / "docs"


def make_dirty_docx():
    """Simulates a real-world SOP written without Word heading styles."""
    doc = Document()

    def bold_line(text):
        p = doc.add_paragraph()
        p.add_run(text).bold = True

    def body(text):
        doc.add_paragraph(text)

    bold_line("SOP-QC-010 OOS（检验结果超标）调查")
    body("版本：2.0    生效日期：2026-05-01    起草部门：质量控制（QC）")
    bold_line("1. 目的")
    body("规定检验结果超出标准（Out of Specification, OOS）时的调查流程，确保在放行决定前查明原因，防止将不合格产品误判为合格或将合格产品误判为不合格。")
    bold_line("2. 范围")
    body("适用于QC实验室所有放行检验、稳定性试验及中间控制检验中出现的OOS结果。趋势异常（OOT）参照本SOP执行但可简化记录。")
    bold_line("3. 职责")
    body("检验员：发现OOS后立即停止后续操作并报告主管，保留全部原始样品、溶液和玻璃器皿。")
    body("QC主管：组织第一阶段实验室调查。")
    body("QA：批准调查结论；涉及产品质量影响时按SOP-QA-001登记偏差。")
    bold_line("4. 调查流程")
    bold_line("4.1 第一阶段Ia：实验室初步调查")
    body("QC主管应在OOS发现后3个工作日内完成初步调查，使用核对清单确认是否存在明显实验室错误（计算错误、稀释错误、仪器故障、标准品失效等）。")
    body("确认为明显实验室错误的，记录错误原因，原始结果作废，重新检验一次并以复验结果报告。")
    bold_line("4.2 第一阶段Ib：实验室扩展调查")
    body("初步调查未发现明显错误时，启动扩展调查，应在OOS发现后20个工作日内完成。可包括对留样的复测：由原检验员和第二名检验员各复测一次，复测方案须经QA事先批准，禁止无方案的重复检验（testing into compliance）。")
    bold_line("4.3 第二阶段：全面调查")
    body("实验室调查不能归因于实验室错误时，扩展至生产环节的全面调查，由QA牵头，按SOP-QA-001登记偏差并开展根本原因分析；必要时按SOP-QA-002开立CAPA。")
    bold_line("4.4 结果处理")
    body("确认的OOS结果涉及已放行批次时，QA应在1个工作日内通知质量受权人，评估是否启动召回程序。")
    bold_line("5. 记录")
    body("OOS调查表、复测方案与结果、调查报告按WI-QA-004保存至少10年。")
    bold_line("6. 参考文件")
    body("SOP-QA-001 偏差管理；SOP-QA-002 CAPA管理；WI-QA-004 文件与记录管理；FDA OOS Guidance (2022)。")

    out = DOCS / "SOP-QC-010_OOS调查.docx"
    doc.save(out)
    return out


def make_pdf():
    """A training-management SOP as a native PDF with numbered sections."""
    lines = [
        "SOP-TR-011 GxP Training Management",
        "Meridian Pharma K.K.   Version 1.3   Effective Date: 2026-06-01",
        "",
        "1. Purpose",
        "This SOP defines requirements for GxP training assignment, delivery,",
        "documentation, and qualification of personnel at Meridian Pharma K.K.",
        "",
        "2. Scope",
        "Applies to all employees and contractors performing GxP activities.",
        "Non-GxP professional development is out of scope.",
        "",
        "3. Training Requirements",
        "3.1 New Hire Induction",
        "All new GxP staff shall complete GxP induction training within 30 calendar",
        "days of start date, and always before independently executing any GxP task.",
        "Until then, tasks shall be performed under documented supervision.",
        "3.2 Annual GMP Refresher",
        "All GxP staff shall complete an annual GMP refresher of at least 8 hours",
        "per calendar year, including data integrity and good documentation practice.",
        "3.3 Role-Based Curricula",
        "Each role shall have a curriculum defined in the qualification matrix.",
        "The qualification matrix shall be reviewed and re-approved annually by",
        "department heads and QA.",
        "",
        "4. Training Delivery and Assessment",
        "4.1 Methods",
        "Training may be delivered as read-and-understand, instructor-led, or",
        "e-learning. Instructor-led GxP training requires a knowledge assessment",
        "with a passing score of at least 80 percent.",
        "4.2 SOP Revision Training",
        "Training on revised documents shall be completed before the effective",
        "date, consistent with WI-QA-004 Section 5.3 (minimum 10 business days",
        "between approval and effective date).",
        "",
        "5. Records",
        "Training records shall be maintained in the validated LMS and retained",
        "for at least 10 years per WI-QA-004. Paper certificates shall be scanned",
        "into the LMS within 5 business days of training completion.",
        "",
        "6. Overdue Training",
        "Training overdue by more than 30 calendar days shall be reported to the",
        "department head; overdue safety-critical training shall suspend the",
        "individual's authorization for the affected task and shall be recorded",
        "as a deviation per SOP-QA-001.",
        "",
        "7. References",
        "SOP-QA-001 Deviation Management; WI-QA-004 Document Control and Records",
        "Management; EU GMP Chapter 2; 21 CFR 211.25.",
    ]
    pdf = fitz.open()
    page = pdf.new_page()
    y, per_page = 72, 46
    count = 0
    for ln in lines:
        if count >= per_page:
            page = pdf.new_page()
            y, count = 72, 0
        page.insert_text((72, y), ln, fontsize=10, fontname="helv")
        y += 15
        count += 1
    out = DOCS / "SOP-TR-011_Training_Management.pdf"
    pdf.save(out)
    pdf.close()
    return out


def make_xlsx():
    """2026 deviation log."""
    wb = Workbook()
    ws = wb.active
    ws.title = "偏差台账2026"
    ws.append(["偏差编号", "记录日期", "部门", "分级", "标题", "状态", "关闭日期", "关联CAPA"])
    rows = [
        ["DEV-2026-001", "2026-01-08", "QC", "Minor", "天平期间核查记录漏签名", "已关闭", "2026-01-20", ""],
        ["DEV-2026-002", "2026-01-15", "PV", "Major", "文献检索周期延迟1周执行", "已关闭", "2026-02-10", "CAPA-2026-003"],
        ["DEV-2026-005", "2026-02-02", "生产", "Minor", "批记录页码装订顺序错误", "已关闭", "2026-02-12", ""],
        ["DEV-2026-008", "2026-02-19", "QA", "Major", "变更CCR-125未完成培训即生效", "已关闭", "2026-03-30", "CAPA-2026-007"],
        ["DEV-2026-011", "2026-03-05", "QC", "Major", "稳定性箱温度超标4小时未报警", "已关闭", "2026-04-15", "CAPA-2026-009"],
        ["DEV-2026-014", "2026-03-28", "PV", "Critical", "安全数据库计划外停机8小时，期间2例ICSR无法录入", "已关闭", "2026-05-06", "CAPA-2026-012"],
        ["DEV-2026-017", "2026-04-14", "IT", "Minor", "eDMS周备份日志未及时归档", "已关闭", "2026-04-28", ""],
        ["DEV-2026-019", "2026-05-09", "QC", "Minor", "对照品领用登记延迟2天", "已关闭", "2026-05-20", ""],
        ["DEV-2026-021", "2026-05-27", "PV", "Critical", "1例严重非预期病例超15天时限迟报PMDA", "调查中", "", "CAPA-2026-015"],
        ["DEV-2026-023", "2026-06-11", "生产", "Major", "洁净区压差记录连续3天未复核", "调查中", "", ""],
        ["DEV-2026-025", "2026-06-30", "QA", "Minor", "内审报告分发清单遗漏一名接收人", "已关闭", "2026-07-08", ""],
        ["DEV-2026-026", "2026-07-06", "QC", "Major", "HPLC审计追踪季度审查逾期12天", "调查中", "", ""],
    ]
    for r in rows:
        ws.append(r)
    out = DOCS / "QA-LOG-2026_偏差台账.xlsx"
    wb.save(out)
    return out


if __name__ == "__main__":
    for f in (make_dirty_docx(), make_pdf(), make_xlsx()):
        print("created", f.name)
