#!/usr/bin/env python3
"""Build formal SOP .docx files from markdown sources in src/ into docs/."""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

SRC = Path(__file__).parent / "src"
OUT = Path(__file__).parent / "docs"

BOLD_RE = re.compile(r"\*\*(.+?)\*\*")


def parse_frontmatter(text):
    meta, body = {}, text
    if text.startswith("---"):
        _, fm, body = text.split("---", 2)
        for line in fm.strip().splitlines():
            k, _, v = line.partition(":")
            meta[k.strip()] = v.strip().strip('"')
    return meta, body.strip()


def add_runs(paragraph, text):
    """Render **bold** spans as bold runs."""
    pos = 0
    for m in BOLD_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        paragraph.add_run(m.group(1)).bold = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def add_md_table(doc, lines):
    rows = [[c.strip() for c in ln.strip().strip("|").split("|")] for ln in lines]
    rows = [r for r in rows if not all(set(c) <= {"-", " ", ":"} for c in r)]
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            p = table.rows[i].cells[j].paragraphs[0]
            add_runs(p, cell)
            if i == 0:
                for run in p.runs:
                    run.bold = True
    return table


def build_doc(md_path):
    meta, body = parse_frontmatter(md_path.read_text(encoding="utf-8"))
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10.5)
    for lvl in range(1, 4):
        h = doc.styles[f"Heading {lvl}"]
        h.font.name = "Arial"
        h.font.color.rgb = RGBColor(0, 0, 0)

    header_p = doc.sections[0].header.paragraphs[0]
    header_p.text = (f"{meta['company']}    {meta['doc_no']} v{meta['version']}    "
                     "CONFIDENTIAL – Fictional document for RAG demo")
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.runs[0].font.size = Pt(8)

    title = doc.add_heading(f"{meta['doc_no']}  {meta['title']}", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)

    info = doc.add_table(rows=4, cols=2)
    info.style = "Table Grid"
    for i, (k, v) in enumerate([("Document No.", meta["doc_no"]),
                                ("Version", meta["version"]),
                                ("Effective Date", meta["effective_date"]),
                                ("Owner", meta["owner"])]):
        info.rows[i].cells[0].paragraphs[0].add_run(k).bold = True
        info.rows[i].cells[1].paragraphs[0].add_run(v)
    doc.add_paragraph()

    lines = body.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line:
            i += 1
            continue
        if line.startswith("|"):
            tbl = [line]
            while i + 1 < len(lines) and lines[i + 1].strip().startswith("|"):
                i += 1
                tbl.append(lines[i])
            add_md_table(doc, tbl)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, line[2:])
        else:
            p = doc.add_paragraph()
            add_runs(p, line)
        i += 1

    out_path = OUT / f"{meta['doc_no']}.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    OUT.mkdir(exist_ok=True)
    for md in sorted(SRC.glob("*.md")):
        print("built", build_doc(md))
